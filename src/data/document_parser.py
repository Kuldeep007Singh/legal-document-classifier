# Code Path: src/data/document_parser.py

import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import io
import re
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
import chardet

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parser for various document formats (PDF, DOCX, TXT)"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.doc']
    
    def parse_document(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse a document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        logger.info(f"Parsing document: {file_path}")
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
            else:
                raise ValueError(f"Parser not implemented for: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return {
                'text': '',
                'metadata': {
                    'file_path': str(file_path),
                    'file_size': file_path.stat().st_size,
                    'error': str(e),
                    'pages': 0,
                    'word_count': 0
                }
            }
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF document using multiple methods for robustness"""
        text = ""
        metadata = {
            'file_path': str(file_path),
            'file_size': file_path.stat().st_size,
            'pages': 0,
            'extraction_method': 'unknown'
        }
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    text = '\n'.join(text_parts)
                    metadata['pages'] = len(pdf.pages)
                    metadata['extraction_method'] = 'pdfplumber'
                    
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {str(e)}")
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    
                    if text_parts:
                        text = '\n'.join(text_parts)
                        metadata['pages'] = len(pdf_reader.pages)
                        metadata['extraction_method'] = 'pypdf2'
                        
            except Exception as e2:
                logger.error(f"PyPDF2 also failed for {file_path}: {str(e2)}")
                metadata['error'] = f"Both extraction methods failed: {str(e)}, {str(e2)}"
        
        # Clean and process the extracted text
        if text:
            text = self._clean_extracted_text(text)
            metadata['word_count'] = len(text.split())
            metadata['char_count'] = len(text)
        
        return {
            'text': text,
            'metadata': metadata
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = '\n'.join(paragraphs)
            text = self._clean_extracted_text(text)
            
            metadata = {
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'word_count': len(text.split()),
                'char_count': len(text),
                'extraction_method': 'python-docx'
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise
    
    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse TXT document with encoding detection"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                text = file.read()
            
            text = self._clean_extracted_text(text)
            
            metadata = {
                'file_path': str(file_path),
                'file_size': file_path.stat().st_size,
                'encoding': encoding,
                'encoding_confidence': encoding_result['confidence'],
                'word_count': len(text.split()),
                'char_count': len(text),
                'extraction_method': 'text_file'
            }
            
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            raise
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page breaks and form feeds
        text = re.sub(r'[\f\r]+', '\n', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def batch_parse_documents(self, file_paths: List[Path], 
                            max_workers: int = 4) -> List[Dict[str, Any]]:
        """Parse multiple documents in parallel"""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from tqdm import tqdm
        
        results = []
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all parsing tasks
            future_to_path = {
                executor.submit(self.parse_document, path): path 
                for path in file_paths
            }
            
            # Collect results with progress bar
            for future in tqdm(as_completed(future_to_path), 
                             total=len(file_paths), 
                             desc="Parsing documents"):
                try:
                    result = future.result()
                    results.append(result)
                except Exception as e:
                    file_path = future_to_path[future]
                    logger.error(f"Failed to parse {file_path}: {str(e)}")
                    results.append({
                        'text': '',
                        'metadata': {
                            'file_path': str(file_path),
                            'error': str(e),
                            'word_count': 0
                        }
                    })
        
        return results
    
    def extract_document_sections(self, text: str) -> Dict[str, str]:
        """Extract common legal document sections"""
        sections = {}
        
        # Common legal document section patterns
        section_patterns = {
            'title': r'(?i)^(.{1,100}?)(?:\n|$)',
            'parties': r'(?i)(parties?|between|agreement.*between)(.*?)(?=\n\n|\n[A-Z]|$)',
            'recitals': r'(?i)(whereas|recitals?)(.*?)(?=\n\n|\nnow therefore|$)',
            'definitions': r'(?i)(definitions?|defined terms?)(.*?)(?=\n\n|\n[0-9]+\.|$)',
            'terms': r'(?i)(terms and conditions|agreement|obligations)(.*?)(?=\n\n|$)',
            'termination': r'(?i)(termination|expir)(.*?)(?=\n\n|$)',
            'governing_law': r'(?i)(governing law|jurisdiction|applicable law)(.*?)(?=\n\n|$)',
            'signatures': r'(?i)(signature|executed|signed)(.*?)'
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
            if match:
                sections[section_name] = match.group().strip()
        
        return sections

class LegalDocumentValidator:
    """Validate legal document content and format"""
    
    def __init__(self):
        self.legal_keywords = [
            'agreement', 'contract', 'party', 'parties', 'whereas', 
            'therefore', 'consideration', 'breach', 'terminate',
            'governing law', 'jurisdiction', 'arbitration', 'dispute',
            'confidential', 'proprietary', 'intellectual property',
            'liability', 'damages', 'indemnify', 'warranty'
        ]
    
    def is_legal_document(self, text: str, threshold: float = 0.1) -> bool:
        """Check if text appears to be a legal document"""
        if not text or len(text) < 100:
            return False
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in self.legal_keywords 
                          if keyword in text_lower)
        
        keyword_density = keyword_count / len(self.legal_keywords)
        
        return keyword_density >= threshold
    
    def validate_document_structure(self, text: str) -> Dict[str, bool]:
        """Validate common legal document structure elements"""
        validations = {
            'has_title': bool(re.search(r'^.{10,100}', text[:200], re.MULTILINE,)),
            'has_parties': bool(re.search(r'(?i)(parties?|between)', text)),
            'has_date': bool(re.search(r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b|\b\w+\s+\d{1,2},?\s+\d{4}\b', text)),
            'has_signature_block': bool(re.search(r'(?i)(signature|signed|executed)', text)),
            'reasonable_length': 100 <= len(text) <= 100000,
            'has_legal_language': self.is_legal_document(text)
        }
        
        return validations

if __name__ == "__main__":
    # Example usage
    parser = DocumentParser()
    validator = LegalDocumentValidator()
    
    # Test with a sample file
    sample_file = Path("data/raw/sample_contract.pdf")
    if sample_file.exists():
        result = parser.parse_document(sample_file)
        print("Extracted text length:", len(result['text']))
        print("Metadata:", result['metadata'])
        
        # Validate the document
        is_legal = validator.is_legal_document(result['text'])
        print("Is legal document:", is_legal)
        
        structure_validation = validator.validate_document_structure(result['text'])
        print("Structure validation:", structure_validation)
    else:
        print(f"Sample file not found: {sample_file}")